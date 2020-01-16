"""
Read word document texts for patients with
Epilepsy, read drugs table, remove sensitive data,
pseudo-anonymise and store keys of personal data in json file
in encrypted environment.

Need to already have created new directory.

This contains all the functions called by main_docx_preprocess

Ali Alim-Marvasti (c) Jan-Feb 2019
"""
# below only works for docx
# so first convert .doc to .docx via .ps1
import docx
import os.path
import os
import re
try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile
import io
import uuid
import json

def args_for_loop(*args):
    """
    factor function used by epilepsy_docx
    makes a list of two numbers
    """
    
    list = []
    for arg in args:
        list.append(arg)
    return list


def update_txt_docx(pt_txt, pt_docx_list, p, clean, print_p_by_p=False):
    """
    This function is a factor function called by epilepsy_docx.
    
    1. pt_txt and pt_docx_list are updated paragraph by paragraph.
    2. p is the paragraph contents from docx Document class.
    3. clean option strips the text
    4. print_p_by_p is an option to print out the text as it is being read.
    
    """
    
    pt_txt = pt_txt + '\n' + p.text
    pt_docx_list.append(p.text)
    
    if print_p_by_p:
        print(p.text)

    if clean:
        pt_txt = pt_txt.strip()

    return pt_txt, pt_docx_list


def save_as_txt(path_to_doc, pt_txt, save_path):
    """
    save as text file in chosen already created directory
    """

    pt_txt = pt_txt.replace('\u03a8', 'Psych')
    pt_txt = pt_txt.replace('\u2192', '>')
    pt_txt = pt_txt.replace('\u03c8', '>')
    pt_txt = pt_txt.replace('\u25ba', '>')
    pt_txt = pt_txt.replace('\uf0e0', 'X')
    pt_txt = pt_txt.replace('\u03bc', 'mu')
    pt_txt = pt_txt.replace('\u2264', '<=')
    pt_txt = pt_txt.replace('\u2248', '~')
    pt_txt = pt_txt.replace('\u03a6', 'PHI')
    pt_txt = pt_txt.replace('\u2265', '>=')
    pt_txt = pt_txt.replace('\u2193', 'down_arrow')
    pt_txt = pt_txt.replace('\u2191', 'up_arrow')
    pt_txt = pt_txt.replace('\u206d', 'up_arrow')
    
    # otherwise gives charmap codec error
    
    try:
        save_filename = path_to_doc.split("\\")[-1].replace(
            '.docx', '')+".txt"  # takes name of docx file
        complete_filename = os.path.join(save_path, save_filename)
        with open(complete_filename, "w") as f:
            print(pt_txt, file=f)
    except UnicodeEncodeError:
        with open(complete_filename, "w", errors='backslashreplace') as f:
            f.write(pt_txt)

def epilepsy_docx_to_txt(path_to_doc, *paragraphs, read_tables=False, clean=False,
                  print_p_by_p=False, docx_to_txt_save_path="L:\\word_docs\\epilepsy_docx_to_txt\\"):
    """
    Returns pt_txt as text.
    Returns pt_docx_list as list of paragraphs.
    Returns pt_meds_list as list of the table - needs further cleanup.
    can also print the text of .docx file.

    path_to_doc includes full folders and filename and extension.

    paragraphs: optional,
      reads between two paragraph numbers (inclusive).
      if not specified, or too many numbers specified, reads all paragraphs.
      if only one number specified, reads from paragraph 0
        to the number specified.

    read_tables will read the tables as paragraphs, if True.

    clean if default false will print(pt_txt) in same format as seen in jupyter
    if true will print(pt_txt) all new lines together with empty lines.
    I prefer clean=True

    print_p_by_p will print paragraph by paragraph as reading the docx file
    """
    document = docx.Document(path_to_doc)

    pt_txt = ''  # initialise text string
    pt_docx_list = []  # initialise patient's docx list
    pt_meds_list = []  # initialise patient meds list
    pt_meds_dict = {}

    if paragraphs:
        para_list = args_for_loop(*paragraphs)

    if len(paragraphs) == 2:
        for n, p in enumerate(document.paragraphs):
            if n >= para_list[0] and n <= para_list[1]:
                pt_txt, pt_docx_list = update_txt_docx(pt_txt, pt_docx_list, p,
                                                       clean, print_p_by_p)

    elif len(paragraphs) == 1:
        for n, p in enumerate(document.paragraphs):
            if n >= 0 and n <= para_list[0]:
                pt_txt, pt_docx_list = update_txt_docx(pt_txt, pt_docx_list, p,
                                                       clean, print_p_by_p)

    else:
        for p in document.paragraphs:
            pt_txt, pt_docx_list = update_txt_docx(pt_txt, pt_docx_list, p,
                                                   clean, print_p_by_p)

    meds_table_headings = ['current rx', 'max dose', 'previous rx', 'stopped/', 'comments', '', ' ']
    if read_tables:
        for t in document.tables:
            for row in t.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text not in meds_table_headings:
                            pt_meds_list.append(paragraph.text)
                pt_meds_dict[row] = pt_meds_list
                        # print(paragraph.text)
        # print (pt_meds_list)

    # this next part cleans to remove empty characters and spaces
    if clean:
        pt_txt = pt_txt.replace('\t', ' ')
        pt_txt = pt_txt.strip()
        
    pt_docx_list = [i.lower() for i in pt_docx_list]
    pt_docx_list = [o for o in pt_docx_list if o != '' and
                    o != '\t' and o != ' ']

    pt_meds_list_clean_and_lower = [
        itm.lower() for itm in pt_meds_list if itm != '' and itm != '\t']
    pt_meds_list_clean_phenytoin = [
        s.replace('  ', '') for s in pt_meds_list_clean_and_lower]
    pt_meds_list_clean_phenytoin = [med.replace(
        'phenytoin ', 'phenytoin') for med in pt_meds_list_clean_phenytoin]
    pt_meds_list = pt_meds_list_clean_phenytoin  # simpler rename

    save_as_txt(path_to_doc, pt_txt, docx_to_txt_save_path)

    return pt_txt, pt_docx_list, pt_meds_dict


def epilepsy_docx_xml_to_txt(path, n_xml, docx_xml_to_txt_save_path="L:\\word_docs\\epilepsy_docx_xml_to_txt\\"):
    """
    Take the path of a docx file as argument, return the text in unicode.
    Run this if epilepsy_docx() isn't able to read the name.
    This should automatically read tables anyway.
    """

    WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    PARA = WORD_NAMESPACE + 'p'
    TEXT = WORD_NAMESPACE + 't'

    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)

    paragraphs = []
    for paragraph in tree.getiterator(PARA):
        texts = [n.text for n in paragraph.getiterator(TEXT) if n.text]
        if texts:
            paragraphs.append(''.join(texts))

    pt_txt_xml = '\n\n'.join(paragraphs)
    save_as_txt(path, pt_txt_xml, docx_xml_to_txt_save_path)

    n_xml += 1
    return pt_txt_xml, n_xml


def name_pattern_regex(name_pattern, pt_txt, fn_index=0, sn_index=-1, IndError=False):
    """
    FACTORISING FUNCTION
    regex search for names
    """
    name_search = re.search(name_pattern, pt_txt)
    name = name_search.group()
    name_list = name.split()

    if IndexError:
        name = name.replace('Name:', '')
        name_list = name.split()

    return name_list


def pt_txt_replace (firstname, surname, pt_txt,
                    redact_message_fname = 'XXXfirstnameXXX',
                    redact_message_sname = 'XXXsurnameXXX'):
    """
    FACTORISING FUNCTION
    replaces firstname and surname in pt_txt with default redact messages.
    """

    pt_txt_fnamefilter = pt_txt.replace(firstname, redact_message_fname)
    pt_txt_sfnamefilter = pt_txt_fnamefilter.replace(surname, redact_message_sname)
    
    names = [firstname, surname]
    
    return pt_txt_sfnamefilter, names


def anonymise_name_txt(pt_txt, path_to_doc, xml=False):
    """
    finds first and surnames using regex and replaces them with
    redact messages XXXfirstnameXXX and XXXsurnameXXX respectively.

    >pt_txt is a string
    >path_to_doc is full filename path including extension
    >xml option's regex is required when text was read by epilepsy_docx_xml.
        this xml option can also be utilised with normal docx conversion.

    Three different regex methods utilised in this function to avoid pitfalls:
    1. list of common template words which are not names
    2. regex variations on common words
    3. negative lookahead regex
    """
    not_pt_names = ['Age', 'age', 'DOB', 'Grand', 'Round', 'Telemetry', 'Meeting', 'Name', 'Name:', ':', 'NAME'
                    'AED', 'Rx', 'XX', 'xx', 'Antecedent', 'Hx', 'telemetry', 'meeting', 'summary', 'ED',
                    'Namexx', 'NamexxAge', 'Hxxx', 'Hx', 'xxx', 'Video', 'video', 'QSA', 'QSB', 'QSC', 'QSD',
                    'Early', 'Development', 'Date', 'TELEMETRY', 'MEETING', 'QS', 'Aged', 'FH', 'No',
                    'Discharge', 'DISCHARGE', 'Summary', 'SUMMARY', 'Family', 'FAMILY', 'No', 'FC', 'ITU'
                    'Referred', 'JD', 'Dr', 'DATE'
                    ]
    regex_notnames = re.compile(r"(Hosp).?")
    regex_notnames_2 = re.compile(r"\s*:\s*")
    NLA = r"(?!Grand|Round|Telemetry|TELEMETRY|MEETING|Meeting|Date|Course|Current|Attacks|Previous|EEG|Imaging|Onset|Risk|Factors|xx|AED|Rx|Hxxx)"
    # NLB = r"(?<![QSABCDEMCVP])"  #  negative lookbehind
    # NLB = r"(?<!Telemetry)"


    if xml:
    # new regex to find names for DOCX - files that don't have names may catch incorrect names
    # some have Surname|SURNAME, Firstname - but must ensure name is there i.e. xml=TRUE
        try:    
            name_pattern = r"%s([A-Z]{1,}[a-z]*,?\s[A-Z]{1}[a-zA-Z]*)"%(NLA) #(?(?=regex)then|else)
            #  negative lookahead to exclude these words
            #  Surname|SURNAME, Firstname
            name_list = name_pattern_regex(name_pattern, pt_txt)
            
            if name_list[0].strip() not in not_pt_names and not regex_notnames.search(name_list[0]) and not regex_notnames_2.search(name_list[0]):
                surname = name_list[0].strip(',')
                surname = surname.strip()
                # clean up any name that picked up suffixes due to xml reader
                surname = re.sub(r"QSA|QSB|QSC|QSD|Telemetry", "", surname)
                
            else: raise AttributeError
            if name_list[-1].strip() not in not_pt_names and not regex_notnames.search(name_list[-1]) and not regex_notnames_2.search(name_list[-1]):
                firstname = name_list[-1]
                firstname = firstname.strip()
                # clean up any name that picked up suffixes due to xml reader
                firstname = re.sub(r"QSA|QSB|QSC|QSD|Telemetry", "", firstname)
                
            else: raise AttributeError

    

            pt_txt_sfnamefilter, names = pt_txt_replace (firstname, surname, pt_txt)

            return pt_txt_sfnamefilter, names                    
        
        except AttributeError:
            raise AttributeError
            # Firstname|FIRSTNAME SURNAME            
            # try:
            #     name_pattern = r"([A-Z]{1}[a-z]+)|([A-Z]{1,}),?\s?\t?([A-Z]+)"  
            #     name_list = name_pattern_regex(name_pattern, pt_txt)
            #     if name_list[0] not in not_pt_names and not regex_notnames.search(name_list[0]) and not regex_notnames_2.search(name_list[0]):
            #         firstname = name_list[0]
            #     elif name_list[1] not in not_pt_names and not regex_notnames.search(name_list[1]) and not regex_notnames_2.search(name_list[1]):
            #         firstname = name_list[1]
            #     else: raise AttributeError

            #     if name_list[-1] not in not_pt_names and not regex_notnames.search(name_list[-1]) and not regex_notnames_2.search(name_list[-1]):
            #         surname = name_list[-1]
            #     else: raise AttributeError

            #     surname2 = re.sub(r"QSA|QSB|QSC|Telemetry|\d", "", surname)
            #     surname2 = surname2.strip()
            #     firstname2 = re.sub(r"QSA|QSB|QSC|Telemetry|\d", "", firstname)
            #     firstname2 = firstname2.strip()
                
            #     pt_txt_sfnamefilter, names = pt_txt_replace (firstname2, surname2, pt_txt)

            #     return pt_txt_sfnamefilter, names
            
            # except IndexError:
            #     raise AttributeError


    else: # if not xml
        try:
            name_pattern = r"((Name|NAME)[\s\t:]*[\s\t]*)([a-zA-Z]+\s+[a-zA-Z]+[\s-]?[a-zA-Z]+)"    
            name_list = name_pattern_regex(name_pattern, pt_txt)

            if name_list[0].strip() not in not_pt_names and not regex_notnames.search(name_list[0]) and not regex_notnames_2.search(name_list[0]):
                firstname = name_list[0]
            elif name_list[1].strip() not in not_pt_names and not regex_notnames.search(name_list[1]) and not regex_notnames_2.search(name_list[1]):
                firstname = name_list[1]
            else:
                firstname = name_list[2]
            if name_list[-1].strip() not in not_pt_names and not regex_notnames.search(name_list[-1]) and not regex_notnames_2.search(name_list[-1]):
                surname = name_list[-1]
            else:
                surname = name_list[name_list.index(firstname) + 1]  # add 1 to the index of the list from firstname

            pt_txt_sfnamefilter, names = pt_txt_replace(firstname, surname, pt_txt)
        
            return pt_txt_sfnamefilter, names

        except IndexError:
            name_list = name_pattern_regex(name_pattern, pt_txt, IndError=True)
            if name_list[0].strip() not in not_pt_names:
                firstname = name_list[0]
            else: raise AttributeError
            if name_list[1].strip() not in not_pt_names:
                surname = name_list[1]
            else: raise AttributeError
            pt_txt_sfnamefilter, names = pt_txt_replace(firstname, surname, pt_txt)

            return pt_txt_sfnamefilter, names

        except AttributeError:  # maybe name is Ali O'Marvasti
            try:
                name_pattern = r"Name[\s\t:]+\w+\s+[A-Z’]+[\s-]?\w+"
                name_list = name_pattern_regex(name_pattern, pt_txt)
                firstname = name_list[1]

                if name_list[-1].strip() not in not_pt_names and not regex_notnames.search(name_list[-1]):
                    surname = name_list[-1]
                elif name_list[2].strip() not in not_pt_names and not regex_notnames.search(name_list[2]):
                    surname = name_list[2]
                else: raise AttributeError

                pt_txt_sfnamefilter, names = pt_txt_replace(firstname, surname, pt_txt) 

                return pt_txt_sfnamefilter, names

            except AttributeError:  # some docx have "surname: SMITH firstname: Harry ..."
                name_pattern = r"surname[\s\t:\n]*[a-z]+\s"
                name_list = name_pattern_regex(name_pattern, pt_txt.lower())

                if name_list[-1].strip() not in not_pt_names and not regex_notnames.search(name_list[-1]):
                    surname = name_list[-1]
                else:
                    print("check this surname out in file {}",format(path_to_doc))
                    raise AttributeError

                name_pattern2 = r"(?<!r)(name)[\s\t:\n]*[a-z]+\s"  # negative look behind
                name_list2 = name_pattern_regex(name_pattern, pt_txt.lower())

                if name_list2[-1].strip() not in not_pt_names and not regex_notnames.search(name_list2[-1]):
                    firstname = name_list2[-1]
                else:
                    print("check this firstname out in file {}",format(path_to_doc))
                    raise AttributeError

                pt_txt_sfnamefilter, names = pt_txt_replace(firstname, surname, pt_txt) 
                # print ("\nused firstname and surname as specified by the file contents {}".format(path_to_doc))
                # print ("firstname is {} and surname is {}".format(firstname, surname))
                return pt_txt_sfnamefilter, names

            

def anonymise_DOB_txt(pt_txt, n_DOB_anon, xml=False):
    """
    finds DOB using regex and replaces with
    pseudoanon DOB e.g. "XXX anonymised DOB = 31/4/63"

    >pt_txt is a string.
    >n_DOB_anon counts number of successfull DOB redactions.

    main_docx_preprocess uses two forms to keep a tally of DOB's needing xml to read.
    """

    try:  # DD/MM/YY(YY) or DD.MM.YY(YY) or DD - MM - YY(YY)
        DOB_pattern = r"D\.?O\.?B\.?[\s\t:]{0,3}([0-9]{1,2}[\s/\.-]?\s?[0-9]{1,2}[\s/\.-]?\s?[0-9]{2,4})"
        DOB_match = re.search(DOB_pattern, pt_txt)
        DOB = DOB_match.group().split()[-1]

    # DDrd FEB YY(YY) or DDth-Feb-YY or . or / or combinations
    except AttributeError:
        try:
            DOB_pattern = r"DOB[\s\t:]+([0-9]{1,2}[st]?[th]?[nd]?[rd]?\s?/?\.?-?\s?[0-9]{0,2}\w*\s?/?\.?-?\s?[0-9]{2,4})"
            DOB_match = re.search(DOB_pattern, pt_txt)
            DOB = DOB_match.group().split()[-1]
        except AttributeError:
            try:
                DOB_pattern = r"DOB[\s\t:]?([0-9]{1,2}[\s/\.-]{1}[0-9]{1,2}[\s/\.-]{1}[0-9]{2,4})"
                DOB_match = re.search(DOB_pattern, pt_txt)
                DOB = DOB_match.group().split()[-1]
            
            except AttributeError:
                try:
                    DOB_pattern = r"(DATE).?(OF).?(BIRTH).?[\s\t:]*([0-9]{1,2}[\s/\.-]?\s?[0-9]{1,2}[\s/\.-]?\s?[0-9]{2,4})"
                    DOB_match = re.search(DOB_pattern, pt_txt)
                    DOB = DOB_match.group().split()[-1]

                except AttributeError:  # finds DOB following hosp number: QDS12345|12345678 dd/mm/yy using positive lookbehind to avoid capture
                    try:
                        DOB_pattern = r"(?<=[A-Z]{3}[0-9]{5}|[0-9]{8})\s?\t?([0-9]{1,2}[\s/\.-]{1}[0-9]{1,2}[\s/\.-]{1}[0-9]{2,4})"
                        DOB_match = re.search(DOB_pattern, pt_txt)
                        DOB = DOB_match.group().split()[-1]
                    
                    except AttributeError:  # this uses findall instead of search and uses the second one as first is date of meeting
                    # note does not catch dd/mm/yyyy when on its own new line  
                        DOB_pattern = r"[0-9]{1,2}/[0-9]{1,2}/[0-9]{2,4}"
                        DOB_findall = re.findall(DOB_pattern, pt_txt)
                        if len(DOB_findall)>1:
                            DOB = DOB_findall[1]
                        else:
                            DOB = DOB_findall

    # sometimes DOB = "DOB:21/1/64" so only keep numbers and / . or -
    # e.g. potential error comes from xml reading DOB where no space between : and DOB
    
    DOB = re.sub("[^0-9/\.-]", "", DOB)  # if 25/Feb/19, then this is removed: 25/19

    # if len((re.compile(r"/?\.?-?").split(DOB))[0].strip()) > 2:  
    #     # if it has somehow accidentally taken DOB as 1234513/01/88
    #     # then keep last two digits(here type are characters) from day
    #     DOB_dd = re.compile(r"/?\.?-?").split(DOB)[0][-2] + re.compile(r"/?\.?-?").split(DOB)[0][-1]
    #     try:
    #         #DOB_mm = re.compile(r"/?\.?-?").split(DOB)[1]
    #         DOB_yy = re.compile(r"/?\.?-?").split(DOB)[-1]
    #         DOB = DOB_dd + '/' + DOB_yy
    #     except ValueError:
    #         pass

    DOB_actual = DOB

    # turn DOB into 3 integers and alter
    # alter DOB so can still trace pt if required
    DOB_digits = list(
        map(lambda x: int(x)+1, (re.compile(r"/?\.?-?").split(DOB))))
    # turn this back to string
    DOB_str = [str(x) for x in DOB_digits]

    if xml:
        try:
            DOB_anon = '\nXXXanonymised DOB=' + \
                DOB_str[0]+'/'+DOB_str[1]+'/'+DOB_str[2]+'XXX' + '\n'
            n_DOB_anon += 1 
            return DOB_anon, n_DOB_anon, DOB_actual

        except IndexError:  # when the month is spelled and not a number i.e. now only 2 items in list
            DOB_anon = '\nXXXanonymised DOB=' + \
                str(DOB_str) +'XXX' + '\n'
            n_DOB_anon += 1 
            return DOB_anon, n_DOB_anon, DOB_actual

    else:
        # and use this message:
        try:
            DOB_anon = 'XXXanonymised DOB=' + \
                DOB_str[0]+'/'+DOB_str[1]+'/'+DOB_str[2] +'XXX'
            pt_txt_DOBfilter = pt_txt.replace(DOB, DOB_anon)
            n_DOB_anon += 1
            return pt_txt_DOBfilter, n_DOB_anon, DOB_actual

        except IndexError:  # when the month is spelled and not a number i.e. now only 2 items in list
            DOB_anon = 'XXXanonymised DOB=' + \
                DOB_str[0]+'/'+DOB_str[1] +'XXX'
            pt_txt_DOBfilter = pt_txt.replace(DOB, DOB_anon)
            n_DOB_anon += 1
            return pt_txt_DOBfilter, n_DOB_anon, DOB_actual

def anon_hosp_no(pt_txt, path_to_doc, uuid_no, n_uuid, n_uuid_name_of_doc):
    """
    Find and replace hospital number (=MRN) with uuid.

    >pt_txt is a string. 
    >path_to_doc is the full pathname and file name with ext.
    >uuid_no is the pseudononymised file number - main_docx passes this argument
    >n_uuid counts number of successful hosp no anonymisations

    >n_uuid_name_of_doc:
        sometimes the document doesn't contain hosp no,
        this function uses MRN in name of file and uses it instead. 
        This argument counts the number of MRNs redacted via this method.

    main_docx_preprocess stores sensitive data as dictionary in .json file.
    (MRN is the actual hospital number innerkey)
    i.e.:
    psuedo_anon_dict: {"MRN": MRN, "Names": [names], "DOB": DOB_actual}

    returns pt_txt without mrn and also the above keys 
    replacing MRN with uuid_no_message
    """
#     rd = random.Random()
#     rd.seed(uuid_rd_seed)  # make it reproducible
#     uuid_no = str(uuid.UUID(int=rd.getrandbits(128)))  # redact_hosp_no_message
    mrn_error_message = False

    try:
        
        # MRN_pattern = r"[QSABCDEMCVP]{0,3}[\s]?(M/)?\d{5,8}((?=\d{2}[\.-/])|(?!\d))"  # exclude X as used in XXX redact messages
        
        # # QSA QSB QSC QSD MCV MVP E
        # # negative lookahead NLA: (?!\d) to avoid catching 2002 telemetry date and appending first part of mrn to synthesise made up mrn.
        # # i.e. Telemetry date 2002 12345678 will catch 12345678 rather than 20021234
        # # however, some files (see pytest complex DOB) are like 1234567810/01/1985 and this messes up to: 34567810 as MRN
        # # so try instead: (?(?=regex)then|else) or (?(?=condition)(then1|then2|then3)|(else1|else2|else3))
        # # (?=\d{1,2}[\.-/])



# # trying #1 to fix QSD1234507 where 07 is DOB dd
#         MRN_pattern = r"([QSABCDMCVP]{3}[\s]?\d{5})|((M/)[\s]?\d{6})|(E[\s]?\d{7})|([\s]?\d{8})" 
#         # fixes error but now again the other error where 2002 date is taken as first four MRN: ie need NLA

# trying #2
        MRN_pattern = r"([QSABCDMCVP]{3}\s?\d{5})|((M/)[\s]?\d{6})|(E[\s]?\d{7})|([\s]?\d{8})((?=\d{2}[\.-/])|(?!\d))" 
        # this works! finally. it does multiple regexes as | ORs and if it doesn't find the first one then tries second
        # the final one has a PLF to find two digits before DOB otherwise it reads the MRN until end of digits. 

        MRN_search = re.search(MRN_pattern, pt_txt.upper())

        MRN = MRN_search.group()
        MRN = MRN.split()
        MRN = MRN[-1]

        uuid_no_message = 'XXX pseudo_anon_dict_DOCX ' + str(uuid_no) + ' XXX'
        pt_txt = pt_txt.upper().replace(MRN, (uuid_no_message))  # .upper otherwise won't replace!
        n_uuid += 1

    except AttributeError:
        try:  # try getting hosp number from document name
            name_of_doc = path_to_doc.split('\\')[-1]

            MRN_pattern = r"[A-Za-z]{0,3}[\s]?\d{5,8}(?!\d)"
            MRN_search = re.search(MRN_pattern, name_of_doc)

            MRN = MRN_search.group()
            MRN = MRN.split()
            MRN = MRN[-1]

            # here need to add the hospital name to the contents of the pt_txt string as it was extracted from document name:
            uuid_no_message = 'XXX pseudo_anon_dict_DOCX ' + str(uuid_no) + ' XXX'
            pt_txt = uuid_no_message + pt_txt
            n_uuid_name_of_doc += 1

        except AttributeError:
            try:  # U/AB1234 hosp no style | MV 52301
                MRN_pattern = r"([A-WYZ]/[A-WYZ]{2}[0-9]{4})|(MV\s\d{5})"
                MRN_search = re.search(MRN_pattern, pt_txt)

                MRN = MRN_search.group()
                #MRN = MRN.split()  # this will change "MV 12345" to two groups"MV" and "12345"
                #MRN = MRN[-1]

                uuid_no_message = 'XXX pseudo_anon_dict_DOCX ' + str(uuid_no) + ' XXX'
                pt_txt = pt_txt.replace(MRN, (uuid_no_message))
                n_uuid += 1
                
            
            except AttributeError:  # admit defeat
                mrn_error_message = True
                MRN = "XXX MRN_ERROR XXX"
                return pt_txt, MRN, n_uuid, n_uuid_name_of_doc, mrn_error_message

    except:
        print("major uncaught error in anon_hosp_no function for \t\t{}\n".format(
            path_to_doc))
        MRN = "XXX MRN_ERROR XXX"
        return pt_txt, MRN, n_uuid, n_uuid_name_of_doc, mrn_error_message

    return pt_txt, MRN, n_uuid, n_uuid_name_of_doc, mrn_error_message
